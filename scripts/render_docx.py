#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PM Master · Markdown -> DOCX 渲染
-------------------------------------------------
双轨产出：Markdown 为单一事实源，本脚本把它渲染为正式 Word 文档。
优先使用 pandoc；若不可用，则用 python-docx 做基础渲染
（支持 #~###### 标题、段落、有序/无序列表、表格、粗体 **x**）。

用法：
  python3 render_docx.py <input.md> [--out output.docx]
"""
import os
import sys
import argparse
import subprocess
import shutil
import re


def _add_code(doc, body, size=9):
    """以等宽字体把纯文本作为代码块写入（用于 mermaid 源码回退）。"""
    from docx.shared import Pt
    p = doc.add_paragraph()
    run = p.add_run(body)
    run.font.name = 'Courier New'
    run.font.size = Pt(size)


def _try_mermaid_png(src):
    """若系统装有 mermaid CLI（mmdc），把 mermaid 源码渲染为 PNG 并返回路径；
    否则返回 None（交由调用方回退为代码块）。失败也不抛错。"""
    mmdc = shutil.which('mmdc')
    if not mmdc:
        return None
    try:
        import tempfile
        with tempfile.NamedTemporaryFile('w', suffix='.mmd', delete=False,
                                         encoding='utf-8') as f:
            f.write(src)
            mmd = f.name
        png = mmd[:-4] + '.png'
        subprocess.run([mmdc, '-i', mmd, '-o', png, '-b', 'white'],
                       check=True, capture_output=True, timeout=60)
        if os.path.exists(png):
            return png
    except Exception:
        return None
    return None


def try_pandoc(md_path, out_path):
    try:
        subprocess.run(['pandoc', md_path, '-o', out_path],
                       check=True, capture_output=True)
        return True
    except (FileNotFoundError, subprocess.CalledProcessError):
        return False


def render_with_python_docx(md_path, out_path):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.read().splitlines()

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        # 围栏代码块（``` ... ```）：mermaid 优先渲染成图，否则以代码块输出，避免乱码
        stripped = line.lstrip()
        if stripped.startswith('```'):
            lang = stripped.strip('`').strip().lower()
            fence_body = []
            i += 1
            while i < n and not lines[i].lstrip().startswith('```'):
                fence_body.append(lines[i])
                i += 1
            if i < n:
                i += 1  # 跳过结束围栏
            body = '\n'.join(fence_body)
            if lang == 'mermaid':
                img = _try_mermaid_png(body)
                if img:
                    try:
                        doc.add_picture(img)
                    except Exception:
                        img = None
                if not img:
                    doc.add_paragraph('[Mermaid 图源 · 本文档未渲染为图片]',
                                     style='Intense Quote')
                    _add_code(doc, body, 8)
            else:
                _add_code(doc, body, 9)
            continue
        # 标题
        m = re.match(r'^(#{1,6})\s+(.*)$', line)
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2).strip(), level=min(level, 6))
            i += 1
            continue
        # 表格（连续 | 行）
        if line.strip().startswith('|'):
            tbl_lines = []
            while i < n and lines[i].strip().startswith('|'):
                tbl_lines.append(lines[i])
                i += 1
            # 去掉分隔行（---）
            rows = [r for r in tbl_lines if not re.match(r'^\s*\|?[\s:\-\|]+\|?\s*$', r)]
            if rows:
                cells = [c.strip() for c in rows[0].strip().strip('|').split('|')]
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = 'Light Grid Accent 1'
                hdr = table.rows[0].cells
                for j, c in enumerate(cells):
                    hdr[j].text = _strip_inline(c)
                for r in rows[1:]:
                    rc = [c.strip() for c in r.strip().strip('|').split('|')]
                    cells_row = table.add_row().cells
                    for j, c in enumerate(rc):
                        if j < len(cells_row):
                            cells_row[j].text = _strip_inline(c)
            continue
        # 无序列表
        if re.match(r'^[-*]\s+', line):
            doc.add_paragraph(_strip_inline(re.sub(r'^[-*]\s+', '', line)), style='List Bullet')
            i += 1
            continue
        # 有序列表
        if re.match(r'^\d+\.\s+', line):
            doc.add_paragraph(_strip_inline(re.sub(r'^\d+\.\s+', '', line)), style='List Number')
            i += 1
            continue
        # 普通段落（合并连续行直到空行）
        para = [line]
        i += 1
        while i < n and lines[i].strip() and not lines[i].lstrip().startswith(('#', '|', '-', '*')) and not re.match(r'^\d+\.\s', lines[i]):
            para.append(lines[i])
            i += 1
        doc.add_paragraph(_strip_inline(' '.join(para)))
    doc.save(out_path)


def _strip_inline(text):
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text


def main():
    ap = argparse.ArgumentParser(description="PM Master MD->DOCX")
    ap.add_argument('input')
    ap.add_argument('--out', default=None)
    a = ap.parse_args()
    out = a.out or (os.path.splitext(a.input)[0] + '.docx')
    if try_pandoc(a.input, out):
        print(f"[docx] pandoc 渲染 -> {out}")
    else:
        try:
            render_with_python_docx(a.input, out)
            print(f"[docx] python-docx 渲染 -> {out}")
        except ImportError:
            print("[docx] 需要 python-docx：pip install python-docx", file=sys.stderr)
            sys.exit(2)
    print(f"[docx] 完成: {out}")


if __name__ == '__main__':
    main()
